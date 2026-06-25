import datetime
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

PROJECT = os.path.dirname(os.path.abspath(__file__))
ORIGINAL_DOCX = r"c:\Users\Bsk\Downloads\elektroničko-poslovanje_dokumentacija.docx"
DOWNLOADS = r"c:\Users\Bsk\Downloads"
STANDALONE = os.path.join(DOWNLOADS, "GiftAI_scraping_objasnjenje.docx")
EXTENDED = os.path.join(DOWNLOADS, "elektroničko-poslovanje_dokumentacija_PROSIRENO.docx")

SHOTS = os.path.join(PROJECT, "screenshots")
COMMAND_FILE = os.path.join(PROJECT, "core", "management", "commands", "import_products.py")
ADMIN_FILE = os.path.join(PROJECT, "core", "admin.py")

AUTHORS = "Boško Raguž, Drago Kulaš, Filip Ostojić i Matija Udovičić"
DATE_STR = datetime.date.today().strftime("%d.%m.%Y.")

_stats = {"images": 0, "tables": 0}

with open(COMMAND_FILE, encoding="utf-8") as f:
    CMD_SRC = f.read()
try:
    with open(ADMIN_FILE, encoding="utf-8") as f:
        ADMIN_SRC = f.read()
except OSError:
    ADMIN_SRC = ""


# ---------- helperi ----------
def extract_function(source, func_name, max_lines=None):
    lines = source.splitlines()
    start, indent = None, 0
    for i, line in enumerate(lines):
        if re.match(rf"^(\s*)def {re.escape(func_name)}\b", line):
            start = i
            indent = len(line) - len(line.lstrip())
            break
    if start is None:
        return ""
    out = [lines[start]]
    for line in lines[start + 1:]:
        if line.strip() and (len(line) - len(line.lstrip())) <= indent and re.match(r"^\s*def ", line):
            break
        out.append(line)
    while out and not out[-1].strip():
        out.pop()
    if max_lines:
        out = out[:max_lines]
    return "\n".join(out)


def extract_assignment(source, varname, max_lines=None):
    lines = source.splitlines()
    out, capturing = [], False
    for line in lines:
        if line.startswith(varname):
            capturing = True
        if capturing:
            out.append(line)
            if line.strip() == "}":
                break
    if max_lines:
        out = out[:max_lines]
    return "\n".join(out)


def extract_class(source, class_name):
    lines = source.splitlines()
    start = None
    for i, line in enumerate(lines):
        if re.match(rf"^class {re.escape(class_name)}\b", line):
            start = i
            break
    if start is None:
        return ""
    out = [lines[start]]
    for line in lines[start + 1:]:
        if line.strip() and not line.startswith((" ", "\t")) and not line.startswith("class " + class_name):
            if re.match(r"^\S", line):
                break
        out.append(line)
    while out and not out[-1].strip():
        out.pop()
    return "\n".join(out)


def H(doc, text, level=1):
    style = "Title" if level == 0 else f"Heading {level}"
    try:
        doc.add_paragraph(text, style=style)
    except KeyError:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(max(11, 18 - level * 2))


def P(doc, text):
    return doc.add_paragraph(text)


def bullet(doc, text):
    try:
        doc.add_paragraph(text, style="List Bullet")
    except KeyError:
        doc.add_paragraph("• " + text)


def mono(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    lines = text.split("\n")
    for j, line in enumerate(lines):
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
        if j < len(lines) - 1:
            run.add_break()
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def image(doc, filename, cap):
    path = os.path.join(SHOTS, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption(doc, cap)
        _stats["images"] += 1
    else:
        P(doc, f"[Umetni screenshot: {filename} – {cap}]")


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    for style_name in ("Light Grid Accent 1", "Light List Accent 1", "Table Grid"):
        try:
            t.style = style_name
            break
        except KeyError:
            continue
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    _stats["tables"] += 1
    doc.add_paragraph()
    return t


# ---------- poglavlje ----------
def add_scraping_chapter(doc, with_title_page=False):
    if with_title_page:
        H(doc, "GiftAI – Web scraping", level=0)
        P(doc, "Objašnjenje izrade web scrapinga za prikupljanje proizvoda")
        P(doc, f"Članovi projekta: {AUTHORS}")
        P(doc, f"Datum: {DATE_STR}")
        doc.add_page_break()
    else:
        H(doc, "Web scraping – proširenje baze proizvoda", level=1)

    # 1
    H(doc, "1. Što je web scraping i zašto ga koristimo", level=2)
    P(doc, "Web scraping je automatizirano prikupljanje podataka s web stranica. Umjesto da "
           "čovjek ručno prepisuje proizvode, program šalje HTTP zahtjev, dohvaća sadržaj "
           "stranice te iz njega izdvaja željene podatke (naziv, cijenu, opis, sliku itd.).")
    P(doc, "Aplikacija GiftAI koristi umjetnu inteligenciju (TF-IDF vektorizacija i kosinusna "
           "sličnost) za preporuku darova. Da bi preporuke bile kvalitetne i raznolike, "
           "potreban je velik katalog proizvoda. Zato sam podatke prikupio web scrapingom, "
           "čime je baza narasla s nekoliko ručno unesenih proizvoda na 1192 proizvoda.")
    image(doc, "01_naslovnica.png", "Slika 1: Početni zaslon aplikacije GiftAI s proizvodima i slikama.")

    # 2
    H(doc, "2. Arhitektura rješenja i tijek podataka", level=2)
    P(doc, "Scraping je realiziran kao Django management naredba (datoteka import_products.py), "
           "pa se cijeli postupak pokreće jednom naredbom: python manage.py import_products. "
           "Podaci prolaze kroz pet koraka:")
    bullet(doc, "1. HTTP zahtjev – šaljemo zahtjev prema izvoru (web stranici ili API-ju) sa "
                "zaglavljem User-Agent.")
    bullet(doc, "2. Parsiranje – HTML razlažemo pomoću BeautifulSoupa, a JSON čitamo izravno.")
    bullet(doc, "3. Čišćenje – sirove vrijednosti (npr. cijena s oznakom valute) pretvaramo u "
                "ispravan format (Decimal), rješavamo relativne URL-ove i sl.")
    bullet(doc, "4. Mapiranje kategorija – kategoriju proizvoda preslikavamo u polja 'interesi' "
                "i 'prigoda' koja AI koristi.")
    bullet(doc, "5. Spremanje – proizvod spremamo u bazu (model Product) zajedno s preuzetom "
                "slikom u Django ImageField.")

    # 3
    H(doc, "3. Tehnika 1 – HTML scraping (books.toscrape.com)", level=2)
    P(doc, "Prvi izvor je books.toscrape.com – stranica izrađena upravo za vježbanje scrapinga. "
           "Koristim biblioteke requests (HTTP) i BeautifulSoup (parsiranje HTML-a). Podatke iz "
           "HTML-a izdvajam CSS selektorima:")
    table(doc,
          ["Podatak", "CSS selektor", "Primjer vrijednosti"],
          [["naziv", "h3 a[title]", "A Light in the Attic"],
           ["cijena", ".price_color", "£51.77"],
           ["kategorija", "ul.breadcrumb li a (3.)", "Poetry"],
           ["opis", "#product_description ~ p", "It's hard to imagine..."],
           ["slika", "#product_gallery img", "korica knjige (.jpg)"]])
    P(doc, "Stranica je podijeljena na 50 podstranica (page-1.html ... page-50.html) pa kroz njih "
           "prolazim u petlji (paginacija). Za svaku knjigu zatim otvaram njezinu detaljnu "
           "stranicu kako bih dohvatio kategoriju, opis i sliku korica.")
    image(doc, "04_izvor_books.png", "Slika 2: Stvarna stranica books.toscrape.com koju scrapeamo.")

    # 4
    H(doc, "4. Čišćenje podataka", level=2)
    P(doc, "Cijena se na stranici pojavljuje kao tekst, npr. \"£51.77\". Regularnim izrazom "
           "uklanjam sve osim znamenki i točke te dobivenu vrijednost pretvaram u Decimal "
           "(51.77), prikladan za spremanje u bazu. Relativne adrese slika i poveznica "
           "pretvaram u apsolutne pomoću funkcije urljoin. Cijeli dohvat je omotan u try/except "
           "kako jedna neispravna stranica ne bi srušila cijeli postupak.")

    # 5
    H(doc, "5. Tehnika 2 – REST API (dummyjson.com)", level=2)
    P(doc, "Drugi izvor je dummyjson.com – besplatni javni REST API. Umjesto HTML-a, on vraća "
           "strukturirani JSON, iz kojeg čitam polja: title, price, category, description i "
           "thumbnail (URL slike). Ovaj izvor donosi raznolikost izvan knjiga (kozmetika, "
           "parfemi, pametni telefoni, dodaci, satovi, naočale itd.) sa slikama koje točno "
           "odgovaraju proizvodu.")
    table(doc,
          ["", "HTML scraping (books.toscrape)", "REST API (dummyjson)"],
          [["Format", "HTML stranica", "JSON"],
           ["Parsiranje", "BeautifulSoup + CSS selektori", "izravno čitanje polja"],
           ["Pouzdanost", "ovisi o strukturi HTML-a", "vrlo stabilno"],
           ["Prednost", "klasična tehnika scrapinga", "strukturirani, čisti podaci"]])
    image(doc, "05_izvor_dummyjson.png", "Slika 3: JSON odgovor dummyjson.com API-ja.")

    # 6
    H(doc, "6. Preuzimanje slika", level=2)
    P(doc, "Za svaki proizvod preuzimam pravu sliku (requests) i spremam je u Django ImageField "
           "pomoću ContentFile. Tako je svaka slika lokalno pohranjena i točno odgovara svom "
           "proizvodu (npr. parfem prikazuje parfem, a ne nasumičnu fotografiju).")

    # 7
    H(doc, "7. Sprječavanje duplikata", level=2)
    P(doc, "Pri spremanju koristim metodu get_or_create po nazivu proizvoda, čime se isti "
           "proizvod ne unosi dvaput ako se naredba pokrene više puta.")

    # 8
    H(doc, "8. Etika i pristojnost prema poslužitelju", level=2)
    P(doc, "Namjerno sam odabrao izvore koji su namijenjeni ili dopušteni za scraping. Stranice "
           "poput Amazona izbjegao sam jer njihovi uvjeti korištenja zabranjuju scraping i "
           "agresivno blokiraju botove. Uz to, šaljem zaglavlje User-Agent i između zahtjeva "
           "dodajem kratku pauzu (time.sleep) kako ne bih preopteretio poslužitelj.")

    # 9
    H(doc, "9. Povezivanje s umjetnom inteligencijom", level=2)
    P(doc, "Prikupljeni podaci nemaju polja 'interesi' i 'prigoda' koja AI koristi za preporuke. "
           "Zato sam izradio mapiranje kategorija (CATEGORY_MAP / DUMMY_MAP) koje svaku "
           "kategoriju proizvoda preslikava u odgovarajuće interese i prigode:")
    table(doc,
          ["Kategorija", "Interesi", "Prigoda"],
          [["Beauty", "ljepota, njega", "Valentinovo, majčin dan"],
           ["Laptops", "tehnologija, posao", "diplomiranje, rođendan"],
           ["Mobile Accessories", "tehnologija, gadgeti", "rođendan, Božić"],
           ["Womens Jewellery", "moda, elegancija", "godišnjica, Valentinovo"],
           ["Poetry (knjiga)", "poezija, romantika, čitanje", "Valentinovo"]])
    P(doc, "Bez ovog koraka AI ne bi mogao uključiti prikupljene proizvode u preporuke. Slika "
           "ispod prikazuje preporuku za upit 'ljepota, parfem' – sustav vraća proizvode iz "
           "kategorije ljepote.")
    image(doc, "02_preporuka.png", "Slika 4: Personalizirana preporuka na temelju unesenih interesa.")
    image(doc, "03_detalj.png", "Slika 5: Detalj proizvoda sa slikom i sličnim proizvodima.")

    # 10
    H(doc, "10. Upravljanje proizvodima kroz Django admin", level=2)
    P(doc, "Prikupljene proizvode pregledavam i uređujem kroz Django administracijsko sučelje. "
           "Admin sam proširio (klasa ProductAdmin): list_display prikazuje naziv, kategoriju, "
           "cijenu i prigodu; search_fields omogućuje pretragu; list_filter dodaje filtriranje "
           "po kategoriji i prigodi. Zahvaljujući tome lako je upravljati s 1192 proizvoda.")
    image(doc, "07_admin_proizvodi.png", "Slika 6: Django admin – popis proizvoda s pretragom i filtrima.")
    image(doc, "08_admin_detalj.png", "Slika 7: Django admin – uređivanje pojedinog proizvoda (uključujući sliku).")
    image(doc, "06_backend_api.png", "Slika 8: Django REST API koji aplikaciji isporučuje proizvode.")

    # 11
    H(doc, "11. Pokretanje", level=2)
    P(doc, "Naredba podržava odabir izvora i ograničenja:")
    bullet(doc, "python manage.py import_products --source all  (knjige + dummyjson)")
    bullet(doc, "python manage.py import_products --source books --limit 200")
    bullet(doc, "python manage.py import_products --source dummyjson --flush  (čisti start)")

    # 12
    H(doc, "12. Rezultat", level=2)
    P(doc, "Konačno stanje baze prikazano je u tablici. Svaki proizvod ima ispravnu sliku i "
           "popunjena polja za AI, a cijeli je postupak ponovljiv jednom naredbom.")
    table(doc,
          ["Izvor", "Tehnika", "Broj proizvoda"],
          [["books.toscrape.com", "HTML scraping (requests + BeautifulSoup)", "~999 (knjige)"],
           ["dummyjson.com", "REST API (JSON)", "193"],
           ["UKUPNO", "", "1192 (svi sa slikom)"]])

    # Dodatak - kod
    H(doc, "Dodatak: isječci izvornog koda", level=2)
    P(doc, "Glavna petlja za scraping knjiga (paginacija, parsiranje, spremanje):")
    mono(doc, extract_function(CMD_SRC, "scrape_books"))
    P(doc, "Dohvat detalja knjige (kategorija, opis, korica) s detaljne stranice:")
    mono(doc, extract_function(CMD_SRC, "_fetch_book_detail"))
    P(doc, "Dohvat proizvoda preko REST API-ja (dummyjson):")
    mono(doc, extract_function(CMD_SRC, "scrape_dummyjson", max_lines=30))
    P(doc, "Preuzimanje i spremanje prave slike proizvoda:")
    mono(doc, extract_function(CMD_SRC, "_ensure_image"))
    P(doc, "Čišćenje cijene u Decimal:")
    mono(doc, extract_function(CMD_SRC, "_parse_price"))
    P(doc, "Mapiranje kategorija knjiga u interese i prigode (isječak):")
    mono(doc, extract_assignment(CMD_SRC, "CATEGORY_MAP"))
    if ADMIN_SRC:
        P(doc, "Proširenje Django admina (ProductAdmin):")
        mono(doc, extract_class(ADMIN_SRC, "ProductAdmin"))


def _save(doc, path):
    try:
        doc.save(path)
        print(f"OK: {path} ({os.path.getsize(path)} bajtova)")
        return True
    except PermissionError:
        print(f"ZAKLJUCANO (preskoceno): {path} - zatvori datoteku u Wordu pa ponovi.")
        return False


def main():
    standalone = Document()
    add_scraping_chapter(standalone, with_title_page=True)
    s_ok = _save(standalone, STANDALONE)

    extended = Document(ORIGINAL_DOCX)
    extended.add_page_break()
    add_scraping_chapter(extended, with_title_page=False)
    e_ok = _save(extended, EXTENDED)

    per_doc_images = _stats["images"] // 2 if (s_ok or e_ok) else 0
    per_doc_tables = _stats["tables"] // 2 if (s_ok or e_ok) else 0
    print(f"Ugradjeno slika (po dokumentu): {per_doc_images}, tablica: {per_doc_tables}")
    if not s_ok:
        print("NAPOMENA: standalone nije spremljen jer je otvoren u Wordu.")


if __name__ == "__main__":
    main()
